'''
cv_generator.py  —  SINGLE FILE
────────────────────────────────
Generates a .docx CV from the *new* JSON schema while preserving the
controlled-random design features introduced earlier.

Only minimal edits were made:
    • Keys renamed to match the latest schema
    • Address now assembled from nested dict
    • “Profile” section inserted if present
    • Field names updated in loops (start_date, end_date, result, date_awarded)
    • _add_two_cols() already fixed to skip empty bullets
'''

from __future__ import annotations
import io, os, random
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
import docx                      # needed for OxmlElement


# ── style pools – unchanged ─────────────────────────────────────────────
FONTS     = ['Calibri', 'Cambria', 'Arial', 'Garamond', 'Georgia', 'Verdana']
BULLETS   = ['•', '–', '◦', '▹']
DIVIDERS  = ['─' * 40, '─' * 20 + ' § ' + '─' * 20, '·' * 40, '—' * 40]
COLOURS   = [
    RGBColor(0x00, 0x4C, 0x99), RGBColor(0x4E, 0x9A, 0x06),
    RGBColor(0xAA, 0x00, 0x00), RGBColor(0x2E, 0x34, 0x36)
]

def _rand_style() -> Dict:
    return dict(
        font=random.choice(FONTS), colour=random.choice(COLOURS),
        bullet=random.choice(BULLETS), divider=random.choice(DIVIDERS),
        headings_upper=random.choice([True, False]),
        name_align=random.choice([WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]),
        show_dividers=random.choice([True, False]),
    )

# ── helpers ────────────────────────────────────────────────────────────
def _get(d: Dict, *keys, default=''):
    for k in keys:
        if k in d: return d[k]
    return default

def _fmt_date(span: str | None) -> str:
    if not span: return 'Present'
    try:         return datetime.strptime(span, '%Y-%m').strftime('%b %Y')
    except:      return span

# ── builders ───────────────────────────────────────────────────────────
heading_size = random.randint(15, 18)
def _add_heading(doc: Document, text: str, style: Dict):
    p = doc.add_paragraph()
    r = p.add_run(text.upper() if style['headings_upper'] else text.title())
    r.bold, r.font.size, r.font.color.rgb = True, Pt(heading_size), style['colour']
    fmt = p.paragraph_format
    fmt.space_before, fmt.space_after, fmt.keep_with_next = Pt(8), Pt(4), True

    p_border = docx.oxml.OxmlElement('w:pBdr')
    bottom   = docx.oxml.OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'C0C0C0')
    p_border.append(bottom); p._p.get_or_add_pPr().append(p_border)

def _add_bullets(doc: Document, items: List[str], style: Dict):
    for line in items:
        para = doc.add_paragraph()
        para.add_run(f"{style['bullet']} ").bold = False
        para.add_run(line)
        fmt = para.paragraph_format
        fmt.space_before, fmt.space_after = Pt(0), Pt(1)

def _add_two_cols(doc: Document, items: List[str], style: Dict):
    n_rows = (len(items) + 1) // 2
    tbl    = doc.add_table(rows=n_rows, cols=2)
    for r in range(n_rows):
        for c in range(2):
            idx = r + c * n_rows
            cell = tbl.cell(r, c)
            if idx < len(items):
                p = cell.paragraphs[0]
                p.text = f'{style["bullet"]} {items[idx]}'
                p.paragraph_format.left_indent = Cm(0.2)
            else:
                cell.text = ''

# ── PUBLIC API ─────────────────────────────────────────────────────────
def cv_json_to_docx(payload: Dict) -> bytes:
    sty = _rand_style()
    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = sty['font'], Pt(random.randint(11,12))

    sec = doc.sections[0]
    for m in ('top_margin', 'bottom_margin'): setattr(sec, m, Cm(1.2))
    for m in ('left_margin', 'right_margin'): setattr(sec, m, Cm(1.5))

    # ── HEADER (always first) ─────────────────────────────────────────
    pd = payload.get('personal_details', {})
    first, last = pd.get('first_name', ''), pd.get('last_name', '')
    h = doc.add_heading(f'{first} {last}'.strip(), level=0)
    h.alignment = sty['name_align']; h.runs[0].font.size = Pt(random.randint(35,50))#22
    h.runs[0].font.color.rgb = sty['colour']

    addr = pd.get('address', {})
    address_line = ', '.join(filter(None, [addr.get('line1'), addr.get('city'), addr.get('country')]))
    contact_line = ' ◆ '.join(filter(None, [pd.get('email'), pd.get('phone'), address_line]))
    if contact_line:
        cp = doc.add_paragraph(contact_line)
        cp.alignment = sty['name_align']; cp.runs[0].font.size = Pt(9)

    # ── PROFILE (always second if present) ────────────────────────────
    if (summary := payload.get('profile')):
        _add_heading(doc, 'Profile', sty)
        doc.add_paragraph(summary)

    # ── PREP REMAINING SECTIONS, THEN SHUFFLE ─────────────────────────
    sections = []

    # 1. Employment history ------------------------------------------------
    if (jobs := payload.get('employment_history')):
        def _write_jobs():
            _add_heading(doc, 'Work History', sty)
            for job in jobs:
                head = f'{job.get("position","")} — {job.get("company","")}'
                p = doc.add_paragraph(head); p.paragraph_format.space_after = Pt(0)
                dates = f'{_fmt_date(job.get("start_date"))} – {_fmt_date(job.get("end_date"))}'
                p.add_run(f'  {dates}').italic = True
                _add_bullets(doc, job.get('responsibilities', []), sty)
        sections.append(_write_jobs)

    # 2. Skills ------------------------------------------------------------
    if (skills := payload.get('skills')):
        def _write_skills():
            _add_heading(doc, 'Skills', sty)
            _add_two_cols(doc, skills, sty)
        sections.append(_write_skills)

    # 3. Education ---------------------------------------------------------
    if (edu := payload.get('education_history')):
        def _write_edu():
            _add_heading(doc, 'Education', sty)
            for ed in edu:
                head = f'{ed.get("degree","")}, {ed.get("institution","")}'
                p = doc.add_paragraph(head)
                dates = f'{_fmt_date(ed.get("start_date"))} – {_fmt_date(ed.get("end_date"))}'
                p.add_run(f'  {dates}').italic = True
                if ed.get('result'):
                    doc.add_paragraph(f'Result: {ed["result"]}')
        sections.append(_write_edu)

    # 4. Languages ---------------------------------------------------------
    if (langs := payload.get('language_qualifications')):
        def _write_langs():
            _add_heading(doc, 'Languages', sty)
            lang_str = [f'{l["language"]} ({l["level"]})' for l in langs]
            _add_two_cols(doc, lang_str, sty)
        sections.append(_write_langs)

    # 5. Certifications ----------------------------------------------------
    if (certs := payload.get('certifications')):
        def _write_certs():
            _add_heading(doc, 'Certifications', sty)
            for c in certs:
                line = f'{c["name"]} — {c["issuer"]} ({_fmt_date(c["date_awarded"])})'
                doc.add_paragraph(line)
        sections.append(_write_certs)

    # --- RANDOMISE order of collected sections ---------------------------
    random.shuffle(sections)
    for write in sections:
        write()

    # ── EXPORT ───────────────────────────────────────────────────────────
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── quick test ────────────────────────────────────────────────────────
def _test():
    sample = {
        "personal_details": {
            "first_name": "Ada",
            "last_name": "Lovelace",
            "address": {
            "line1": "12 St James’s Sq",
            "city": "London",
            "country": "UK"
            },
            "phone": "+44 20 7946 0958",
            "email": "ada.lovelace@alumni.london.ac.uk"
        },
        "profile": "Visionary mathematician and the world's first computer programmer. Recognized for her work on Charles Babbage's Analytical Engine and her insights into the future potential of computing. Known for blending analytical rigor with imaginative foresight, she laid the groundwork for modern algorithms.",
        "employment_history": [
            {
            "position": "Mathematical Analyst and Scientific Collaborator",
            "company": "Analytical Engine Laboratory",
            "location": "London, UK",
            "start_date": "1840-01",
            "end_date": "1843-12",
            "responsibilities": [
                "Authored the first published algorithm intended for implementation on a mechanical computing device.",
                "Translated and annotated Menabrea’s paper on the Analytical Engine, expanding it threefold with original commentary.",
                "Worked closely with Charles Babbage to conceptualize the practical applications of computational machines.",
                "Pioneered thinking on the potential of computers to go beyond mere arithmetic."
            ]
            }
        ],
        "education_history": [
            {
            "degree": "Bachelor of Science (BSc)",
            "field": "Mathematics",
            "institution": "University of London",
            "location": "London, UK",
            "start_date": "1831-10",
            "end_date": "1835-07",
            "result": "First-class honours"
            }
        ],
        "language_qualifications": [
            {
            "language": "English",
            "level": "native"
            },
            {
            "language": "French",
            "level": "fluent"
            }
        ],
        "certifications": [
            {
            "name": "Royal Society Fellowship",
            "issuer": "Royal Society",
            "date_awarded": "1848-05"
            }
        ],
        "skills": [
            "Mathematical Analysis",
            "Algorithm Design",
            "Technical Writing"
        ]
        }

    

    data = cv_json_to_docx(sample)
    os.makedirs('output', exist_ok=True)
    with open('output/test_cv.docx', 'wb') as f: f.write(data)
    print('Generated → output/test_cv.docx')

if __name__ == '__main__':
    _test()
