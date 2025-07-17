'''
cv_generator.py  —  SINGLE FILE
────────────────────────────────
Generates a .docx CV from the *new* JSON schema while preserving the
controlled-random design features introduced earlier.

Only minimal edits were made:
    • Keys renamed to match the latest schema
    • Address now assembled from nested dict
    • “Profile” section inserted if present
    • Field names updated in loops (start_date, end_date, result, date_awarded)
    • _add_two_cols() already fixed to skip empty bullets
'''

from __future__ import annotations
import io, os, random
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx                      # needed for OxmlElement


# ── style pools – unchanged ─────────────────────────────────────────────
FONTS     = ['Calibri', 'Cambria', 'Arial', 'Garamond', 'Georgia', 'Verdana']
BULLETS   = ['•', '–', '◦', '▹']
DIVIDERS  = ['─' * 40, '─' * 20 + ' § ' + '─' * 20, '·' * 40, '—' * 40]
COLOURS   = [
    RGBColor(0x00, 0x4C, 0x99), RGBColor(0x4E, 0x9A, 0x06),
    RGBColor(0xAA, 0x00, 0x00), RGBColor(0x2E, 0x34, 0x36)
]

def _rand_style() -> Dict:
    return dict(
        font=random.choice(FONTS), colour=random.choice(COLOURS),
        bullet=random.choice(BULLETS), divider=random.choice(DIVIDERS),
        headings_upper=random.choice([True, False]),
        name_align=random.choice([WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]),
        show_dividers=random.choice([True, False]),
        border=random.choice(['single', 'double', 'dotted', 'dashed', 'dotDash']),
        template=random.randint(0, 4),    # 0-4 choose among five templates
    )

# ── helpers ────────────────────────────────────────────────────────────
def _get(d: Dict, *keys, default=''):
    for k in keys:
        if k in d: return d[k]
    return default

def _fmt_date(span: str | None) -> str:
    if not span: return 'Present'
    try:         return datetime.strptime(span, '%Y-%m').strftime('%b %Y')
    except:      return span

# ── builders ───────────────────────────────────────────────────────────
heading_size = random.randint(15, 18)
def _add_heading(doc: Document, text: str, style: Dict):
    p = doc.add_paragraph()
    r = p.add_run(text.upper() if style['headings_upper'] else text.title())
    r.bold, r.font.size, r.font.color.rgb = True, Pt(heading_size), style['colour']
    fmt = p.paragraph_format
    fmt.space_before, fmt.space_after, fmt.keep_with_next = Pt(8), Pt(4), True
    if style.get('border') not in (None, '', 'none'):
        p_border = docx.oxml.OxmlElement('w:pBdr')
        bottom   = docx.oxml.OxmlElement('w:bottom')
        bottom.set(qn('w:val'), style['border'])
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'C0C0C0')
        p_border.append(bottom)
        p._p.get_or_add_pPr().append(p_border)

    return p   # ← we need the para back for later tweaks


def _add_bullets(doc: Document, items: List[str], style: Dict):
    for line in items:
        para = doc.add_paragraph()
        para.add_run(f"{style['bullet']} ").bold = False
        para.add_run(line)
        fmt = para.paragraph_format
        fmt.space_before, fmt.space_after = Pt(0), Pt(1)

def _add_two_cols(doc: Document, items: List[str], style: Dict):
    n_rows = (len(items) + 1) // 2
    tbl    = doc.add_table(rows=n_rows, cols=2)
    for r in range(n_rows):
        for c in range(2):
            idx = r + c * n_rows
            cell = tbl.cell(r, c)
            if idx < len(items):
                p = cell.paragraphs[0]
                p.text = f'{style["bullet"]} {items[idx]}'
                p.paragraph_format.left_indent = Cm(0.2)
            else:
                cell.text = ''
# ── TIMELINE helpers (template 1) ───────────────────────────────────────
def _tl_dates(rec, start_k='start_date', end_k='end_date'):
    start, end = rec.get(start_k), rec.get(end_k)

    # If neither date is supplied, suppress the whole block
    if not start and not end:
        return ''

    # Normal formatting rules
    start_txt = _fmt_date(start) if start else ''
    end_txt   = _fmt_date(end)   if end   else 'Present'

    # Avoid a leading/trailing dash if one side is missing
    return f'{start_txt} – {end_txt}'.strip(' –')

def _write_timeline(container, title, rows, sty, start_k, end_k, fmt_fn):
    _add_heading(container, title, sty)
    tbl = container.add_table(rows=len(rows), cols=2)
    tbl.columns[0].width, tbl.columns[1].width = Cm(3), Cm(13)
    for i, rec in enumerate(rows):
        tbl.cell(i, 0).text = _tl_dates(rec, start_k, end_k)
        cell = tbl.cell(i, 1)
        fmt_fn(cell.paragraphs[0], rec)

        # ── NEW: add bullets for job responsibilities (if any) ───────────
        if rec.get("responsibilities"):
            for line in rec["responsibilities"]:
                para = cell.add_paragraph()
                para.add_run(f"{sty['bullet']} ")
                para.add_run(line)
                fmt = para.paragraph_format
                fmt.left_indent = Cm(0.2)
                fmt.space_before, fmt.space_after = Pt(0), Pt(1)

def _tl_work(p, rec):
    p.add_run(rec.get('position','')).bold = True
    p.add_run(f' — {rec.get("company","")}')

def _tl_edu(p, rec):
    p.add_run(rec.get('degree','')).bold = True
    p.add_run(f', {rec.get("institution","")}')
    if rec.get("location"):
        p.add_run(f' — {rec["location"]}')

# ── section-writer factory (shared by all templates) ────────────────────
def _section_writers(container, payload, sty):
    w = {}
    # Work
    if (jobs := payload.get('employment_history')):
        def work():
            _add_heading(container, 'Work History', sty)
            for jb in jobs:
                head = f'{jb.get("position","")} — {jb.get("company","")}'
                p = container.add_paragraph(head); p.paragraph_format.space_after = Pt(0)
                p.add_run(f'  {_tl_dates(jb)}').italic = True
                _add_bullets(container, jb.get('responsibilities', []), sty)
        w['work'] = work
    # Education
    if (edu := payload.get('education_history')):
        def edu_w():
            _add_heading(container, 'Education', sty)
            for ed in edu:
                # build "Degree, Institution, Location" but drop blanks
                parts = [ed.get("degree",""), ed.get("institution",""), ed.get("location","")]
                head  = ', '.join(filter(None, parts))
                p = container.add_paragraph(head)
                dates = _tl_dates(ed)
                if dates:                       # only add when something to show
                    p.add_run(f'  {dates}').italic = True
                if ed.get('result'):
                    container.add_paragraph(f'Result: {ed["result"]}')
        w['edu'] = edu_w
    # Skills
    if (skills := payload.get('skills')):
        w['skills'] = lambda: (_add_heading(container,'Skills',sty), _add_two_cols(container,skills,sty))
    # Languages
    if (langs := payload.get('language_qualifications')):
        def langs_w():
            _add_heading(container,'Languages',sty)
            items=[f'{l["language"]} ({l["level"]})' for l in langs]
            _add_two_cols(container,items,sty)
        w['langs'] = langs_w
    # Certs
    if (certs := payload.get('certifications')):
        def certs_w():
            _add_heading(container,'Certifications',sty)
            for c in certs:
                container.add_paragraph(f'{c["name"]} — {c["issuer"]} ({_fmt_date(c["date_awarded"])})')
        w['certs'] = certs_w
    return w

def _rough_line_count(doc, chars_per_line=80):
    """Cheap heuristic: estimate printed lines from total characters."""
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    return max(1, round(total_chars / chars_per_line))

def _ensure_one_page(doc, sec, body_style_name='Normal'):
    # ❶ Figure out how many lines normally fit on page 1
    #    Lines ≈ (usable vertical space / line height)
    usable_cm = (sec.page_height.cm - (sec.top_margin.cm + sec.bottom_margin.cm))
    # 11 pt ≈ 0.39 cm; add 20 % breathing room
    lines_per_page = int(usable_cm / (0.39 * 1.2))

    # ❷ Compare with rough content length
    if _rough_line_count(doc) >= lines_per_page:
        return                      # already ≥ 1 page — do nothing

    # ❸ Inflate: bump font one point & widen margins
    body = doc.styles[body_style_name].font
    body.size = Pt((body.size.pt or 11) + 1)      # +1 pt everywhere

    # grow top & bottom only; left/right unchanged
    sec.top_margin  = Cm(sec.top_margin.cm  + 1)   # +1.2 cm
    sec.bottom_margin = Cm(sec.bottom_margin.cm + 1)

# ── PUBLIC API ─────────────────────────────────────────────────────────
def cv_json_to_docx(payload: Dict, template: int | None = None) -> bytes:
    sty = _rand_style()
    if template is not None:           # caller forces template 0-4
        sty['template'] = max(0, min(4, template))
    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = sty['font'], Pt(random.randint(12,13))

    sec = doc.sections[0]
    for m in ('top_margin', 'bottom_margin'): setattr(sec, m, Cm(1.2))
    for m in ('left_margin', 'right_margin'): setattr(sec, m, Cm(1.5))

    # ── HEADER (always first) ─────────────────────────────────────────
    pd = payload.get('personal_details', {})
    first, last = pd.get('first_name', ''), pd.get('last_name', '')
    # python-docx ≥ 1.0 does **not** create a run automatically
    h = doc.add_heading("", level=0)
    name_text = f"{first} {last}".strip() or "Name"
    run = h.add_run(name_text)            # ensure at least one run
    h.alignment = sty["name_align"]
    run.font.size = Pt(random.randint(40, 50))
    run.font.color.rgb = sty["colour"]

    addr = pd.get('address', {})
    address_line = ', '.join(filter(None, [addr.get('line1'), addr.get('city'), addr.get('country')]))

    contact_line = ' ◆ '.join(filter(None, [pd.get('email'), pd.get('phone'), address_line]))
    if contact_line:
        cp = doc.add_paragraph(contact_line)
        cp.alignment = sty['name_align']; cp.runs[0].font.size = Pt(9)

    # ── PROFILE ───────────────────────────────────────────────────────
    if (summary := payload.get('profile')):
        _add_heading(doc, 'Profile', sty); doc.add_paragraph(summary)

    # ── BODY via templates ────────────────────────────────────────────
    writers = _section_writers(doc, payload, sty)
    t = sty['template']

    if t == 0:  # classic random order
        order = list(writers.values()); random.shuffle(order)
        for w in order: w()

    elif t == 1:  # timeline layout
        if payload.get('employment_history'):
            _write_timeline(doc, 'Work History',
                            payload['employment_history'], sty,
                            'start_date', 'end_date', _tl_work)
        if payload.get('education_history'):
            _write_timeline(doc, 'Education',
                            payload['education_history'], sty,
                            'start_date', 'end_date', _tl_edu)
        for k in ('skills','langs','certs'):
            if k in writers: writers[k]()

    elif t == 2:  # shaded headings, education first
        def shade(par, rgb):
            shd = docx.oxml.OxmlElement('w:shd')
            shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*rgb))
            par._p.get_or_add_pPr().append(shd)
            for run in par.runs: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        original_add_heading = _add_heading
        def add_shaded(doc_, text_, st=sty):
            p = original_add_heading(doc_, text_, st)
            shade(p, st['colour'])
            return p
        globals()['_add_heading'], old = add_shaded, _add_heading
        for k in ('edu','work','skills','langs','certs'):
            if k in writers: writers[k]()
        globals()['_add_heading'] = old

    elif t == 3:  # sidebar
        tbl = doc.add_table(rows=1, cols=2)
        tbl.columns[0].width = Cm(5)
        left, right = tbl.cell(0,0), tbl.cell(0,1)
        left._tc.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w'))))
        for k in ('skills','langs','certs'):
            if k in _section_writers(left, payload, sty): _section_writers(left,payload,sty)[k]()
        for k in ('work','edu'):
            if k in _section_writers(right, payload, sty): _section_writers(right,payload,sty)[k]()

    else:        # t == 4 minimalist: fixed order, no borders
        sty['border']='none'
        for k in ('work','skills','edu','langs','certs'):
            if k in writers: writers[k]()


    # ── EXPORT ───────────────────────────────────────────────────────────
    _ensure_one_page(doc, sec)          # ← NEW: inflate if < 1 page
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── quick test ────────────────────────────────────────────────────────
def _test():
    sample = {
        "personal_details": {
            "first_name": "Ada",
            "last_name": "Lovelace",
            "address": {
            "line1": "12 St James’s Sq",
            "city": "London",
            "country": "UK"
            },
            "phone": "+44 20 7946 0958",
            "email": "ada.lovelace@alumni.london.ac.uk"
        },
        "profile": "Visionary mathematician and the world's first computer programmer. Recognized for her work on Charles Babbage's Analytical Engine and her insights into the future potential of computing. Known for blending analytical rigor with imaginative foresight, she laid the groundwork for modern algorithms.",
        "employment_history": [
            {
            "position": "Mathematical Analyst and Scientific Collaborator",
            "company": "Analytical Engine Laboratory",
            "location": "London, UK",
            "start_date": "1840-01",
            "end_date": "1843-12",
            "responsibilities": [
                "Authored the first published algorithm intended for implementation on a mechanical computing device.",
                "Translated and annotated Menabrea’s paper on the Analytical Engine, expanding it threefold with original commentary.",
                "Worked closely with Charles Babbage to conceptualize the practical applications of computational machines.",
                "Pioneered thinking on the potential of computers to go beyond mere arithmetic."
            ]
            }
        ],
        "education_history": [
            {
            "degree": "Bachelor of Science (BSc)",
            "field": "Mathematics",
            "institution": "University of London",
            "location": "London, UK",
            "start_date": "1831-10",
            "end_date": "1835-07",
            "result": "First-class honours"
            }
        ],
        "language_qualifications": [
            {
            "language": "English",
            "level": "native"
            },
            {
            "language": "French",
            "level": "fluent"
            }
        ],
        "certifications": [
            {
            "name": "Royal Society Fellowship",
            "issuer": "Royal Society",
            "date_awarded": "1848-05"
            }
        ],
        "skills": [
            "Mathematical Analysis",
            "Algorithm Design",
            "Technical Writing"
        ]
        }
    sample = {
                
                "personal_details": {
                    "first_name": "Selamwit",
                    "last_name": "Habtamu",
                    "address": {
                    "line1": "16 Rose Gardens",
                    "city": "120 - Black - African or African British",
                    "country": "United Kingdom"
                    },
                    "phone": "07521069073",
                    "email": "habtamuselamwit@gmail.com"
                },
                "profile": "Motivated and detail-oriented aspiring hospitality professional pursuing a BA (Hons) Hospitality Management degree at UWL. Experienced in maintaining cleanliness and hygiene standards, with strong teamwork and time management skills. Committed to delivering quality service in dynamic environments.",
                "employment_history": [
                    {
                    "position": "Cleaner",
                    "company": "",
                    "location": "United Kingdom",
                    "start_date": "2021-03",
                    "end_date": "2025-08",
                    "responsibilities": [
                        "Maintain cleanliness and hygiene of assigned areas including floors, surfaces, and restrooms.",
                        "Dispose of waste and recyclables according to health and safety guidelines.",
                        "Ensure cleaning equipment is used and stored properly to avoid hazards.",
                        "Report maintenance issues and supply needs promptly to management."
                    ]
                    }
                ],
                "education_history": [
                    {
                    "degree": "Secondary Education",
                    "field": "General Studies",
                    "institution": "Unknown Secondary School",
                    "location": "Eritrea",
                    "start_date": "",
                    "end_date": "",
                    "result": ""
                    }
                ],
                "language_qualifications": [
                    {
                    "language": "English",
                    "level": "Intermediate"
                    },
                    {
                    "language": "Tigrinya",
                    "level": "Native"
                    }
                ],
                "certifications": [],
                "skills": [
                    "Time Management",
                    "Teamwork",
                    "Communication",
                    "Attention to Detail",
                    "Use of Cleaning Equipment",
                    "Sanitation and Hygiene",
                    "Health and Safety Compliance",
                    "Waste Management"
                ]
                
            }
        
    

    

    os.makedirs('output', exist_ok=True)
    for t in range(6):
        if t == 5:
            data = cv_json_to_docx(sample)
        else:                     # 0-4 → all templates
            data = cv_json_to_docx(sample, template=t)
        filename = f'output/test_cv_t{t}.docx'
        with open(filename, 'wb') as f:
            f.write(data)
        print(f'Generated → {filename}')

if __name__ == '__main__':
    _test()
