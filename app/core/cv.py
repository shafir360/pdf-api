from __future__ import annotations

import io
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches


# ───────────────────────── helpers ───────────────────────── #

def _add_heading(doc: Document, text: str) -> None:
    """Section heading: UPPER-CASE, underline, 12 pt."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    p.style = "Normal"
    # underline (bottom border)
    p_format = p.paragraph_format
    p_format.space_before = Pt(12)
    p_format.space_after = Pt(4)
    p_format.left_indent = Cm(0)
    p_format.keep_together = True
    p_format.keep_with_next = True

    p_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_border.append(bottom)
    p._p.get_or_add_pPr().append(p_border)


def _add_bullets(doc: Document, items: List[str]) -> None:
    for line in items:
        para = doc.add_paragraph(line, style="List Bullet")
        para.paragraph_format.space_after = Pt(2)


def _add_two_column_list(doc: Document, items: List[str]) -> None:
    half = (len(items) + 1) // 2
    t = doc.add_table(rows=half, cols=2)
    t.autofit = True
    for r in range(half):
        for c in range(2):
            try:
                t.cell(r, c).text = items[r + c * half]
            except IndexError:
                t.cell(r, c).text = ""
            t.cell(r, c).paragraphs[0].style = "List Bullet"


def _fmt_date(span: str) -> str:
    try:
        return datetime.strptime(span, "%Y-%m").strftime("%b %Y")
    except Exception:
        return span or "Present"


# ───────────────────────── public API ────────────────────── #

def cv_json_to_docx(payload: Dict) -> bytes:
    """Create a `.docx` CV with magazine-style formatting."""
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Reduce margins for a compact one-pager
    s = doc.sections[0]
    for m in ("top_margin", "bottom_margin"):
        setattr(s, m, Cm(1.2))
    for m in ("left_margin", "right_margin"):
        setattr(s, m, Cm(1.5))

    # ─── Header ────────────────────────────────────────────
    personal = payload.get("personal", {})
    name = (personal.get("first_name", "") + " " + personal.get("last_name", "")).strip()
    if personal.get("full_caps", False):
        name = name.upper()

    h = doc.add_heading(name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.size = Pt(22)

    title = personal.get("title")
    if title:
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    contact = payload.get("contact", {})
    contact_line = " ◆ ".join(filter(None, [
        contact.get("email"), contact.get("phone"),
        contact.get("linkedin"), contact.get("github"),
        contact.get("address")
    ]))
    if contact_line:
        p = doc.add_paragraph(contact_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)

    # ─── Profile / Summary ────────────────────────────────
    if personal.get("summary"):
        _add_heading(doc, "Professional Summary")
        doc.add_paragraph(personal["summary"])

    # ─── Experience ───────────────────────────────────────
    if (exp := payload.get("experience")):
        _add_heading(doc, "Work History")
        for job in exp:
            headline = f"{job.get('position','')} — {job.get('company','')}"
            para = doc.add_paragraph(headline, style="Heading 3")
            para.paragraph_format.space_after = Pt(0)
            dates = f"{_fmt_date(job.get('start',''))} – {_fmt_date(job.get('end','Present'))}"
            para.add_run(f"  {dates}").italic = True
            _add_bullets(doc, job.get("responsibilities", []))

    # ─── Education ────────────────────────────────────────
    if (edu := payload.get("education")):
        _add_heading(doc, "Education")
        for ed in edu:
            headline = f"{ed.get('degree','')}, {ed.get('institution','')}"
            para = doc.add_paragraph(headline, style="Heading 3")
            dates = f"{_fmt_date(ed.get('start',''))} – {_fmt_date(ed.get('end',''))}"
            para.add_run(f"  {dates}").italic = True
            if ed.get("grade"):
                doc.add_paragraph(f"Grade: {ed['grade']}")

    # ─── Skills (two-column) ──────────────────────────────
    if (skills := payload.get("skills")):
        _add_heading(doc, "Skills")
        _add_two_column_list(doc, skills)

    # ─── Projects ─────────────────────────────────────────
    if (projs := payload.get("projects")):
        _add_heading(doc, "Projects")
        for pr in projs:
            head = pr.get("name", "")
            tech = ", ".join(pr.get("tech", []))
            para = doc.add_paragraph(head, style="Heading 3")
            if tech:
                para.add_run(f"  ({tech})").italic = True
            if pr.get("description"):
                doc.add_paragraph(pr["description"])

    # ─── Misc sections in same style ──────────────────────
    for key, nice in [("certifications", "Certifications"),
                      ("languages", "Languages"),
                      ("interests", "Interests")]:
        if (items := payload.get(key)):
            _add_heading(doc, nice)
            if isinstance(items, list):
                _add_two_column_list(doc, items)
            else:
                doc.add_paragraph(str(items))

    # ─── References line (optional) ───────────────────────
    if payload.get("references", True):
        doc.add_paragraph("References available upon request.").italic = True

    # ─── Return bytes ─────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
